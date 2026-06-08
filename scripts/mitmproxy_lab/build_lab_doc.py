"""Construire le guide de lab d'observation du trafic Claude Opus 4.7 par proxy.

Ce livrable accompagne le document llm_stateless_architecture.docx et
fournit une mise en oeuvre operationnelle pas-a-pas pour Linux et
Windows. Il genere :

  - un document Word avec quatre schemas Mermaid (architecture, modes
    proxy, trust store, scenarios) et la procedure complete ;
  - trois scripts auxiliaires (setup_linux.sh, setup_windows.ps1,
    analyze_traffic.py) embarques dans les annexes du Word ;
  - un zip des assets pour partage.

Toutes les variables d'environnement Claude Code et les commandes de
trust store ont ete verifiees sur la documentation publique 2026
(docs.claude.com/en/docs/claude-code/network-config, doc mitmproxy 12).

Audience : expert cybersecurite, professeur M2 IA/cyber, analyste SOC
voulant instrumenter des agents LLM en lab.
"""
from __future__ import annotations

import json
import logging
import os
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

WORK_DIR: Path = Path(__file__).parent.resolve()
DIAGRAMS_DIR: Path = WORK_DIR / "diagrams_lab"
OUTPUT_DOCX: Path = WORK_DIR / "llm_lab_proxy_opus47.docx"

CHROME_BIN: str = os.environ.get(
    "PUPPETEER_EXECUTABLE_PATH",
    "/root/.cache/puppeteer/chrome/linux-148.0.7778.167/chrome-linux64/chrome",
)

MERMAID_CONFIG: dict = {
    "theme": "default",
    "themeVariables": {
        "background": "#ffffff",
        "primaryColor": "#ffffff",
        "primaryTextColor": "#222222",
        "primaryBorderColor": "#444444",
        "lineColor": "#444444",
        "secondaryColor": "#f3f3f3",
        "tertiaryColor": "#ffffff",
        "fontFamily": "Arial, sans-serif",
        "fontSize": "15px",
    },
}

PUPPETEER_CONFIG: dict = {"args": ["--no-sandbox", "--disable-setuid-sandbox"]}

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("lab_doc_builder")


# ---------------------------------------------------------------------------
# Diagram dataclass
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class Diagram:
    """Metadata for a generated diagram."""

    key: str
    title: str
    caption: str
    source: str = ""
    width_cm: float = 16.5


# ---------------------------------------------------------------------------
# Mermaid diagrams (UTF-8 with French accents)
# ---------------------------------------------------------------------------


MERMAID_SOURCES: list[Diagram] = [
    Diagram(
        key="lab01_architecture",
        title="Schéma 1 - Architecture du lab",
        caption=(
            "Vue d'ensemble du dispositif d'observation. Le poste de "
            "travail (Linux ou Windows) héberge à la fois le client "
            "(Claude Code CLI, extension VS Code Continue ou Cline, "
            "SDK Python) et le proxy mitmproxy en écoute locale. Tout "
            "le trafic HTTPS vers api.anthropic.com transite par "
            "mitmproxy, est déchiffré grâce à la CA mitmproxy "
            "installée dans le trust store du poste, journalisé par "
            "l'addon Python, puis rechiffré et acheminé vers Anthropic."
        ),
        source="""flowchart LR
    subgraph POSTE [Poste de travail Linux ou Windows]
        direction TB
        CLI[Claude Code CLI<br/>ou extension VS Code]:::client
        TRUST[Trust store du système<br/>contient la CA mitmproxy]:::trust
        PROXY[mitmproxy<br/>localhost:8888]:::proxy
        ADDON[addon Python<br/>llm_traffic_logger.py]:::addon
        JSONL[(traffic.jsonl)]:::store
    end

    CLI -->|HTTPS_PROXY ou<br/>ANTHROPIC_BASE_URL| PROXY
    CLI -.->|vérifie chaîne TLS| TRUST
    PROXY -->|HTTPS reformé| API[api.anthropic.com<br/>Claude Opus 4.7]:::anthropic
    API -->|SSE chunks| PROXY
    PROXY -->|SSE pass-through| CLI

    PROXY --> ADDON
    ADDON --> JSONL

    classDef client fill:#e3eaf3,stroke:#3a5e89,color:#1a2a3a
    classDef trust fill:#fff4d6,stroke:#b58900,color:#5c4400
    classDef proxy fill:#ffe0b2,stroke:#e65100,color:#4e2200
    classDef addon fill:#f4e8ff,stroke:#7e57c2,color:#3d2466
    classDef store fill:#f5f5f5,stroke:#888,color:#333
    classDef anthropic fill:#dff5e1,stroke:#2f7d32,color:#1b3d1f
""",
    ),
    Diagram(
        key="lab02_modes_proxy",
        title="Schéma 2 - Mode regular versus mode reverse",
        caption=(
            "Deux configurations possibles pour intercepter le trafic "
            "Anthropic. Le mode regular (HTTPS_PROXY) est l'approche "
            "universelle, le client garde l'URL api.anthropic.com et "
            "le proxy intercepte la connexion via la variable "
            "d'environnement standard. Le mode reverse (apiBase "
            "override) est plus simple à mettre en oeuvre mais "
            "nécessite que le client accepte une URL personnalisée "
            "comme ANTHROPIC_BASE_URL."
        ),
        source="""flowchart TB
    subgraph REG [Mode regular - HTTPS_PROXY]
        direction LR
        C1[Client]:::client
        P1[mitmproxy<br/>mode regular<br/>port 8888]:::proxy
        A1[api.anthropic.com]:::anthropic
        C1 -->|HTTPS_PROXY=<br/>http localhost 8888<br/>conserve l'URL Anthropic| P1
        P1 -->|CONNECT puis TLS<br/>avec CA mitmproxy| A1
    end

    subgraph REV [Mode reverse - ANTHROPIC_BASE_URL]
        direction LR
        C2[Client]:::client
        P2[mitmproxy<br/>mode reverse<br/>port 8888]:::proxy
        A2[api.anthropic.com]:::anthropic
        C2 -->|ANTHROPIC_BASE_URL=<br/>http localhost 8888<br/>URL réécrite| P2
        P2 -->|forward HTTPS<br/>vers api.anthropic.com| A2
    end

    classDef client fill:#e3eaf3,stroke:#3a5e89,color:#1a2a3a
    classDef proxy fill:#ffe0b2,stroke:#e65100,color:#4e2200
    classDef anthropic fill:#dff5e1,stroke:#2f7d32,color:#1b3d1f
""",
    ),
    Diagram(
        key="lab03_trust_store",
        title="Schéma 3 - Chaîne de confiance de la CA mitmproxy",
        caption=(
            "La CA mitmproxy doit être ajoutée au bon trust store pour "
            "que les clients TLS l'acceptent. Sur Linux, deux niveaux : "
            "le trust store système (utilisé par curl, Python requests, "
            "Node.js depuis la v22) et le store NSS (utilisé par "
            "Firefox et Chromium). Sur Windows, le Certificate Store "
            "couvre la majorité des clients y compris les outils en "
            "ligne de commande. Claude Code v2.1.101 et ultérieur trust "
            "par défaut le trust store OS via CLAUDE_CODE_CERT_STORE = "
            "bundled,system."
        ),
        source="""flowchart TB
    CA[CA mitmproxy<br/>mitmproxy-ca-cert.pem<br/>généré au premier lancement]:::ca

    CA --> LIN[Linux Ubuntu Debian]:::os
    CA --> WIN[Windows 11]:::os

    LIN --> L1[/usr local share ca-certificates<br/>puis update-ca-certificates/]:::store_sys
    LIN --> L2[~/.pki/nssdb<br/>via certutil -A NSS]:::store_nss

    WIN --> W1[Cert:\\LocalMachine\\Root<br/>via certutil -addstore]:::store_sys
    WIN --> W2[Cert:\\CurrentUser\\Root<br/>via Import-Certificate]:::store_sys

    L1 -.->|trust| CL1[curl, wget, Python requests]:::client
    L1 -.->|trust| CL2[Node.js v22+, claude.exe Bun]:::client
    L2 -.->|trust| CL3[Firefox, Chromium snap]:::client

    W1 -.->|trust| CW1[curl.exe, PowerShell Invoke-WebRequest]:::client
    W1 -.->|trust| CW2[Edge, Chrome]:::client
    W1 -.->|trust| CW3[Claude Code CLI, VS Code Electron]:::client

    classDef ca fill:#222,color:#fff,stroke:#222
    classDef os fill:#e3eaf3,stroke:#3a5e89,color:#1a2a3a
    classDef store_sys fill:#fff4d6,stroke:#b58900,color:#5c4400
    classDef store_nss fill:#f4e8ff,stroke:#7e57c2,color:#3d2466
    classDef client fill:#dff5e1,stroke:#2f7d32,color:#1b3d1f
""",
    ),
    Diagram(
        key="lab04_scenarios",
        title="Schéma 4 - Scénarios d'observation programmés",
        caption=(
            "Quatre scénarios à dérouler dans l'ordre pour vérifier le "
            "bon fonctionnement du lab et observer les phénomènes "
            "décrits dans la note technique principale. Chaque scénario "
            "produit des marqueurs spécifiques dans le JSONL qui "
            "permettent de valider l'interception."
        ),
        source="""flowchart TD
    S1[Scénario 1 - Premier message<br/>1 tour, sans cache_control<br/>vérifier message_start, deltas, message_stop]:::scenario
    S2[Scénario 2 - Cache write puis hit<br/>2 requêtes identiques en moins de 5 min<br/>vérifier cache_creation_input_tokens<br/>puis cache_read_input_tokens]:::scenario
    S3[Scénario 3 - Boucle tool_use<br/>demande qui force read_file<br/>vérifier stop_reason tool_use<br/>puis tool_result au tour suivant]:::scenario
    S4[Scénario 4 - Session longue<br/>20 tours conversationnels<br/>vérifier croissance du payload input<br/>et ratio cache hit]:::scenario

    S1 --> S2 --> S3 --> S4

    S1 -.->|sortie attendue| O1[1 entrée JSONL<br/>usage.input_tokens, output_tokens]:::output
    S2 -.->|sortie attendue| O2[2 entrées JSONL<br/>2e contient cache_read_input_tokens > 0]:::output
    S3 -.->|sortie attendue| O3[2 entrées JSONL minimum<br/>1re avec content_block type tool_use<br/>2e avec tool_result en user message]:::output
    S4 -.->|sortie attendue| O4[20 entrées JSONL<br/>analyze_traffic.py calcule la courbe]:::output

    classDef scenario fill:#e3eaf3,stroke:#3a5e89,color:#1a2a3a
    classDef output fill:#dff5e1,stroke:#2f7d32,color:#1b3d1f
""",
    ),
]


# ---------------------------------------------------------------------------
# Mermaid rendering
# ---------------------------------------------------------------------------


def write_supporting_files(target_dir: Path) -> tuple[Path, Path]:
    """Write the Mermaid theme and Puppeteer config files used by mmdc."""
    target_dir.mkdir(parents=True, exist_ok=True)
    mermaid_cfg = target_dir / "mermaid_config.json"
    puppeteer_cfg = target_dir / "puppeteer_config.json"
    mermaid_cfg.write_text(json.dumps(MERMAID_CONFIG, indent=2), encoding="utf-8")
    puppeteer_cfg.write_text(json.dumps(PUPPETEER_CONFIG, indent=2), encoding="utf-8")
    return mermaid_cfg, puppeteer_cfg


def render_mermaid(diagram: Diagram, out_dir: Path) -> Path:
    """Render a single Mermaid diagram to a PNG file via mmdc."""
    out_dir.mkdir(parents=True, exist_ok=True)
    src_path = out_dir / f"{diagram.key}.mmd"
    png_path = out_dir / f"{diagram.key}.png"
    src_path.write_text(diagram.source, encoding="utf-8")
    mermaid_cfg, puppeteer_cfg = write_supporting_files(out_dir)
    env = os.environ.copy()
    env["PUPPETEER_EXECUTABLE_PATH"] = CHROME_BIN
    cmd = [
        "mmdc",
        "-i", str(src_path),
        "-o", str(png_path),
        "-c", str(mermaid_cfg),
        "-p", str(puppeteer_cfg),
        "-b", "white",
        "-w", "1800",
    ]
    logger.info("Rendering %s", diagram.key)
    result = subprocess.run(cmd, env=env, capture_output=True, text=True, check=False)
    if result.returncode != 0 or not png_path.exists():
        logger.error("mmdc failed for %s: %s", diagram.key, result.stderr or result.stdout)
        raise RuntimeError(f"Mermaid rendering failed for {diagram.key}")
    return png_path


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, hex_color: str) -> None:
    """Apply a background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "C62828", size: str = "12") -> None:
    """Apply a colored border on all sides of a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_toc(document: Document) -> None:
    """Insert a placeholder TOC field, refreshed by Word on open."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Mettre à jour cette table dans Word (clic droit, Mettre à jour les champs)."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(el)


def add_code_block(document: Document, code: str, language: str = "bash") -> None:
    """Insert a monospace code block with subtle gray background."""
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.width = Cm(16.5)
    tag_p = cell.paragraphs[0]
    tag_run = tag_p.add_run(language)
    tag_run.bold = True
    tag_run.font.name = "Consolas"
    tag_run.font.size = Pt(8)
    tag_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for line in code.splitlines():
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1C)


def add_caption(document: Document, text: str) -> None:
    """Add a small italic caption under a figure."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_figure(document: Document, image_path: Path, diagram: Diagram) -> None:
    """Insert a diagram (heading, figure, caption)."""
    document.add_heading(diagram.title, level=3)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(diagram.width_cm))
    add_caption(document, diagram.caption)


def add_bullet(document: Document, text: str) -> None:
    """Add a list bullet paragraph."""
    p = document.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def add_callout_box(
    document: Document,
    title: str,
    paragraphs: list[str],
    bg_hex: str = "FDECEA",
    border_hex: str = "C62828",
    title_color: tuple[int, int, int] = (0xC6, 0x28, 0x28),
) -> None:
    """Insert a colored callout box with title and body paragraphs."""
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_hex)
    set_cell_border(cell, color=border_hex, size="12")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.width = Cm(16.5)
    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(*title_color)
    for text in paragraphs:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x3A, 0x1A, 0x1A)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_security_box(document: Document, title: str, paragraphs: list[str]) -> None:
    """Cybersecurity callout (red)."""
    add_callout_box(
        document,
        "Lecture cybersécurité — " + title,
        paragraphs,
        bg_hex="FDECEA",
        border_hex="C62828",
        title_color=(0xC6, 0x28, 0x28),
    )


def add_tip_box(document: Document, title: str, paragraphs: list[str]) -> None:
    """Operational tip callout (green)."""
    add_callout_box(
        document,
        "Astuce opérationnelle — " + title,
        paragraphs,
        bg_hex="E8F5E9",
        border_hex="2F7D32",
        title_color=(0x1B, 0x5E, 0x20),
    )


def configure_styles(document: Document) -> None:
    """Set default fonts and heading appearance."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[level]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0x1A, 0x2A, 0x3A)
        style.font.bold = True


# ---------------------------------------------------------------------------
# Document content - sections
# ---------------------------------------------------------------------------


def _build_cover(doc: Document) -> None:
    """Cover page."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(110)
    run = title.add_run("Mise en oeuvre d'un lab d'observation")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x3A)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("du trafic Claude Opus 4.7 via mitmproxy")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x3A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(20)
    run = subtitle.add_run("Procédure détaillée Linux et Windows")
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(40)
    run = meta.add_run(
        "Note technique de laboratoire. Usage interne et pédagogique. "
        "Audience experte cybersécurité et IA."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _section_1_objectives(doc: Document, diagrams: dict) -> None:
    doc.add_heading("1. Objectifs et architecture du lab", level=1)
    doc.add_paragraph(
        "Cette note décrit la mise en oeuvre d'un dispositif "
        "d'observation du trafic réseau entre un client IDE et le "
        "modèle Claude Opus 4.7 hébergé sur api.anthropic.com. "
        "L'objectif est de capturer, déchiffrer et journaliser les "
        "échanges HTTP en clair applicatif pour analyse pédagogique "
        "ou recherche sécurité."
    )
    doc.add_paragraph(
        "Le dispositif repose sur mitmproxy en mode regular, sur le "
        "même poste que le client, avec la CA mitmproxy installée "
        "dans le trust store du système. L'addon Python "
        "llm_traffic_logger.py (fourni dans la note technique "
        "principale) tape le trafic chunk par chunk sans casser le "
        "streaming SSE et écrit un fichier JSONL exploitable hors "
        "ligne."
    )
    add_figure(doc, *diagrams["lab01_architecture"])

    doc.add_heading("1.1 Ce que le lab permet d'observer", level=2)
    add_bullet(doc, "Le payload exact envoyé par le client (system prompt, tools, messages, model)")
    add_bullet(doc, "Les en-têtes HTTP incluant x-api-key et anthropic-version")
    add_bullet(doc, "Le flux SSE événement par événement (message_start, deltas, message_delta, message_stop)")
    add_bullet(doc, "Les compteurs usage retournés par Anthropic : input_tokens, output_tokens, cache_creation_input_tokens, cache_read_input_tokens")
    add_bullet(doc, "Les blocs tool_use et tool_result au fil de la boucle agentique")
    add_bullet(doc, "Les valeurs stop_reason et leur évolution")
    add_bullet(doc, "Les latences réseau et de génération")


def _section_2_modes(doc: Document, diagrams: dict) -> None:
    doc.add_heading("2. Choix architectural : mode regular ou mode reverse", level=1)
    doc.add_paragraph(
        "Deux configurations sont possibles pour intercepter le "
        "trafic. Le choix entre les deux dépend du client utilisé "
        "et du degré de réalisme souhaité par rapport à une "
        "interception MITM en conditions de production."
    )
    add_figure(doc, *diagrams["lab02_modes_proxy"])

    doc.add_heading("2.1 Mode regular (HTTPS_PROXY)", level=2)
    doc.add_paragraph(
        "Le client conserve l'URL api.anthropic.com et délègue le "
        "transport à mitmproxy via la variable d'environnement "
        "standard HTTPS_PROXY. mitmproxy reçoit la requête CONNECT, "
        "interceptat la session TLS en présentant un certificat "
        "signé par sa CA, puis ré-établit une session TLS vers le "
        "vrai serveur Anthropic. Tous les clients respectueux de "
        "HTTPS_PROXY (curl, Python requests, Node.js, Claude Code "
        "CLI, Continue, Cline) fonctionnent. C'est la configuration "
        "recommandée pour un lab fidèle au comportement de "
        "production."
    )

    doc.add_heading("2.2 Mode reverse (ANTHROPIC_BASE_URL)", level=2)
    doc.add_paragraph(
        "Le client est configuré pour parler directement à "
        "localhost:8888 via la variable ANTHROPIC_BASE_URL. "
        "mitmproxy en mode reverse forwarde vers "
        "https://api.anthropic.com. Aucun déchiffrement TLS n'est "
        "nécessaire côté client (la liaison entre le client et "
        "mitmproxy peut rester en HTTP clair). Plus simple à "
        "mettre en oeuvre, mais ne permet pas d'observer le "
        "comportement d'un client en conditions réseau réelles "
        "avec certificate validation."
    )

    add_tip_box(
        doc,
        "quel mode choisir",
        [
            "Mode regular : pour reproduire fidèlement le "
            "comportement d'un client en production, valider une "
            "politique de TLS inspection, ou former à la détection "
            "d'un MITM réel. Demande l'installation de la CA dans "
            "le trust store. Compatible avec tous les clients.",

            "Mode reverse : pour itérer rapidement sans manipuler "
            "les trust stores, faire des démonstrations courtes, "
            "ou tester un client qui ne respecte pas HTTPS_PROXY. "
            "Compatible avec les clients qui acceptent une URL "
            "d'API personnalisée (ANTHROPIC_BASE_URL pour Claude "
            "Code CLI et SDK officiel, apiBase pour Continue et "
            "Cline). Ne reproduit pas la chaîne TLS.",
        ],
    )


def _section_3_prereq(doc: Document) -> None:
    doc.add_heading("3. Pré-requis", level=1)

    doc.add_heading("3.1 Compte Anthropic et clé API", level=2)
    doc.add_paragraph(
        "Le lab consomme du Claude Opus 4.7 sur api.anthropic.com. "
        "Il est nécessaire de disposer :"
    )
    add_bullet(doc, "D'un compte Anthropic Console avec accès facturé.")
    add_bullet(doc, "D'une clé API au format sk-ant-api03-... générée depuis la Console.")
    add_bullet(doc, "D'un budget de quelques dizaines de centimes pour les scénarios de test (typiquement moins de 1 euro pour une session complète de validation).")

    add_security_box(
        doc,
        "gestion de la clé API",
        [
            "La clé API doit être stockée hors du code. Ne pas la "
            "commiter dans Git, ne pas la mettre dans un fichier de "
            "configuration partagé. Sous Linux, utiliser un fichier "
            "~/.anthropic chmodé à 600 ou un secret manager (pass, "
            "gopass, age, Bitwarden CLI). Sous Windows, utiliser "
            "Windows Credential Manager via cmdkey ou un coffre.",

            "Une clé exfiltrée par un proxy compromis ou par une "
            "extension VS Code malveillante peut être utilisée "
            "jusqu'à révocation manuelle dans la Console. Surveiller "
            "le compteur d'usage sur la Console pendant le lab et "
            "révoquer immédiatement en cas d'anomalie.",
        ],
    )

    doc.add_heading("3.2 Logiciels requis", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Composant"
    hdr[1].text = "Linux"
    hdr[2].text = "Windows 11"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "1A2A3A")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rows = [
        ("Python", "3.10 ou supérieur, apt install python3 pipx",
         "3.11 ou supérieur, winget install Python.Python.3.12"),
        ("mitmproxy", "pipx install mitmproxy (version 12 ou supérieure)",
         "pip install mitmproxy (version 12 ou supérieure)"),
        ("certutil", "apt install libnss3-tools (pour NSS)",
         "Inclus dans Windows par défaut"),
        ("Claude Code CLI", "npm install -g @anthropic-ai/claude-code",
         "npm install -g @anthropic-ai/claude-code"),
        ("VS Code", "snap install code ou .deb officiel",
         "winget install Microsoft.VisualStudioCode"),
        ("Extension Continue", "Marketplace VS Code, dernière version stable",
         "Marketplace VS Code, dernière version stable"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
    doc.add_paragraph()

    doc.add_heading("3.3 Vérifications préalables", level=2)
    doc.add_paragraph(
        "Avant de commencer, valider que le poste peut joindre "
        "api.anthropic.com en direct, sans interférence d'un proxy "
        "d'entreprise ou d'un outil de TLS inspection déjà actif "
        "(Zscaler, Cato Networks, CrowdStrike Falcon)."
    )
    add_code_block(
        doc,
        '''# Linux
curl -sS -o /dev/null -w "%{http_code}\\n" -m 8 https://api.anthropic.com/
openssl s_client -connect api.anthropic.com:443 -servername api.anthropic.com \\
  </dev/null 2>/dev/null | openssl x509 -noout -issuer

# Windows PowerShell
(Invoke-WebRequest -Uri "https://api.anthropic.com/" -Method Head -UseBasicParsing).StatusCode
(New-Object System.Net.Http.HttpClient).GetAsync("https://api.anthropic.com/").Result.StatusCode''',
        language="bash",
    )
    doc.add_paragraph(
        "L'émetteur du certificat doit être une autorité publique "
        "reconnue (par exemple Amazon ou Cloudflare). Si l'émetteur "
        "est une CA d'entreprise, un dispositif de TLS inspection "
        "est déjà en place et viendra interférer avec mitmproxy. "
        "Dans ce cas, exécuter le lab sur une machine personnelle "
        "ou un réseau invité, pas sur le réseau d'entreprise."
    )


def _section_4_linux(doc: Document, diagrams: dict) -> None:
    doc.add_heading("4. Installation et configuration Linux", level=1)
    doc.add_paragraph(
        "Procédure validée sur Ubuntu 22.04 LTS, Ubuntu 24.04 LTS et "
        "Debian 12. Toutes les commandes sont exécutables dans un "
        "shell utilisateur avec accès sudo."
    )

    doc.add_heading("4.1 Installation de mitmproxy", level=2)
    doc.add_paragraph(
        "Préférer pipx à pip pour isoler mitmproxy dans son propre "
        "environnement virtuel et éviter les conflits avec les "
        "packages système."
    )
    add_code_block(
        doc,
        '''sudo apt update
sudo apt install -y python3 python3-pip pipx libnss3-tools
pipx ensurepath
pipx install mitmproxy
# Recharger le PATH dans le shell courant
source ~/.bashrc

# Vérifier la version (12 ou supérieur attendu)
mitmproxy --version | head -n 1''',
        language="bash",
    )

    doc.add_heading("4.2 Génération de la CA mitmproxy", level=2)
    doc.add_paragraph(
        "La CA est générée automatiquement au premier lancement de "
        "n'importe quel outil mitmproxy. Un lancement de 3 secondes "
        "suffit, on peut tuer le processus juste après."
    )
    add_code_block(
        doc,
        '''mitmdump --listen-port 18888 &
sleep 3
kill %1
wait %1 2>/dev/null

# Vérifier que les fichiers existent
ls -la ~/.mitmproxy/
# mitmproxy-ca.pem, mitmproxy-ca-cert.pem, mitmproxy-ca-cert.cer,
# mitmproxy-ca-cert.p12, mitmproxy-dhparam.pem''',
        language="bash",
    )
    add_figure(doc, *diagrams["lab03_trust_store"])

    doc.add_heading("4.3 Installation dans le trust store système", level=2)
    doc.add_paragraph(
        "Le trust store système (couvre curl, wget, Python requests, "
        "Node.js depuis v22, claude.exe) est mis à jour via "
        "update-ca-certificates. La CA doit être renommée en .crt "
        "(extension obligatoire pour update-ca-certificates) et "
        "placée dans /usr/local/share/ca-certificates."
    )
    add_code_block(
        doc,
        '''sudo cp ~/.mitmproxy/mitmproxy-ca-cert.pem \\
    /usr/local/share/ca-certificates/mitmproxy-ca-cert.crt
sudo update-ca-certificates

# Vérifier que la CA est bien intégrée au bundle système
grep -c "mitmproxy" /etc/ssl/certs/ca-certificates.crt
# Doit retourner au moins 1''',
        language="bash",
    )

    doc.add_heading("4.4 Installation dans le store NSS (Firefox, Chromium)", level=2)
    doc.add_paragraph(
        "Firefox et Chromium installé via snap utilisent leur propre "
        "base NSS pour la validation TLS. L'ajout doit être fait "
        "explicitement avec certutil."
    )
    add_code_block(
        doc,
        '''# Initialiser la base NSS si elle n'existe pas encore
mkdir -p ~/.pki/nssdb
[ -f ~/.pki/nssdb/cert9.db ] || certutil -d sql:$HOME/.pki/nssdb -N --empty-password

# Ajouter la CA mitmproxy
certutil -A -d sql:$HOME/.pki/nssdb \\
    -n mitmproxy \\
    -t "C,," \\
    -i ~/.mitmproxy/mitmproxy-ca-cert.pem

# Vérifier
certutil -L -d sql:$HOME/.pki/nssdb | grep mitmproxy''',
        language="bash",
    )

    doc.add_heading("4.5 Variables d'environnement", level=2)
    doc.add_paragraph(
        "Préparer les variables dans ~/.bashrc ou ~/.zshrc. On ne "
        "les active pas systématiquement : pour les laisser "
        "commentées et n'activer que pendant les sessions de lab, "
        "on utilise un sourcing manuel ou une fonction shell."
    )
    add_code_block(
        doc,
        '''# A coller dans ~/.bashrc

# Toujours actif : pointe Node.js et Python requests vers la CA mitmproxy
export NODE_EXTRA_CA_CERTS="$HOME/.mitmproxy/mitmproxy-ca-cert.pem"
export SSL_CERT_FILE="$HOME/.mitmproxy/mitmproxy-ca-cert.pem"

# Cle Anthropic, jamais commitee
export ANTHROPIC_API_KEY="sk-ant-api03-..."

# Fonction shell pour activer ou desactiver le proxy a la demande
llmlab() {
    case "$1" in
        on)
            export HTTPS_PROXY=http://localhost:8888
            export HTTP_PROXY=http://localhost:8888
            echo "LLM lab proxy activated"
            ;;
        off)
            unset HTTPS_PROXY HTTP_PROXY
            echo "LLM lab proxy deactivated"
            ;;
        *)
            echo "Usage: llmlab on|off"
            ;;
    esac
}''',
        language="bash",
    )

    add_tip_box(
        doc,
        "script d'installation automatisé",
        [
            "L'annexe A contient setup_linux.sh qui automatise les "
            "étapes 4.1 à 4.5. Pour l'utiliser :",
            "  chmod +x setup_linux.sh && ./setup_linux.sh",
            "Le script journalise chaque étape et signale les "
            "anomalies (absence de certutil, échec de connectivité "
            "Anthropic).",
        ],
    )


def _section_5_windows(doc: Document) -> None:
    doc.add_heading("5. Installation et configuration Windows 11", level=1)
    doc.add_paragraph(
        "Procédure validée sur Windows 11 23H2 et 24H2. Toutes les "
        "commandes sont exécutables dans PowerShell 7 ou Windows "
        "PowerShell 5.1. L'installation de la CA en mode utilisateur "
        "ne nécessite pas de privilèges administrateur ; "
        "l'installation machine-wide oui."
    )

    doc.add_heading("5.1 Installation de Python", level=2)
    add_code_block(
        doc,
        '''# Verifier la presence de Python
py --version

# Installer si necessaire
winget install --id Python.Python.3.12 --silent

# Recharger la session ou ouvrir un nouveau terminal
$env:PATH = [System.Environment]::GetEnvironmentVariable('PATH','Machine') + ';' + `
            [System.Environment]::GetEnvironmentVariable('PATH','User')

py --version
py -m pip --version''',
        language="powershell",
    )

    doc.add_heading("5.2 Installation de mitmproxy", level=2)
    add_code_block(
        doc,
        '''py -m pip install --upgrade pip
py -m pip install --upgrade mitmproxy

# Verifier la version
mitmdump --version | Select-Object -First 1''',
        language="powershell",
    )

    doc.add_heading("5.3 Génération de la CA mitmproxy", level=2)
    add_code_block(
        doc,
        '''$proc = Start-Process -FilePath mitmdump `
    -ArgumentList "--listen-port","18888" `
    -PassThru -NoNewWindow
Start-Sleep -Seconds 3
Stop-Process -Id $proc.Id -Force -ErrorAction SilentlyContinue

# Verifier que les fichiers existent
Get-ChildItem "$env:USERPROFILE\\.mitmproxy"''',
        language="powershell",
    )

    doc.add_heading("5.4 Import dans le Certificate Store Windows", level=2)
    doc.add_paragraph(
        "Deux niveaux possibles : utilisateur courant (sans "
        "privilèges) ou machine (administrateur). Le niveau machine "
        "est nécessaire si plusieurs comptes utilisent le poste, ou "
        "pour que les services Windows (qui tournent sous "
        "LocalSystem) acceptent la CA."
    )
    add_code_block(
        doc,
        '''# Import dans CurrentUser - sans privileges administrateur
Import-Certificate `
    -FilePath "$env:USERPROFILE\\.mitmproxy\\mitmproxy-ca-cert.cer" `
    -CertStoreLocation Cert:\\CurrentUser\\Root

# Variante avec certutil (-user pour user store)
certutil -user -addstore -f Root `
    "$env:USERPROFILE\\.mitmproxy\\mitmproxy-ca-cert.cer"

# Import machine-wide - necessite PowerShell en administrateur
# Import-Certificate `
#     -FilePath "$env:USERPROFILE\\.mitmproxy\\mitmproxy-ca-cert.cer" `
#     -CertStoreLocation Cert:\\LocalMachine\\Root

# Verifier
Get-ChildItem Cert:\\CurrentUser\\Root | Where-Object Subject -match "mitmproxy"''',
        language="powershell",
    )

    doc.add_heading("5.5 Variables d'environnement utilisateur", level=2)
    doc.add_paragraph(
        "Les variables sont persistées au niveau User scope via "
        "[Environment]::SetEnvironmentVariable. Elles sont "
        "automatiquement disponibles dans toutes les nouvelles "
        "sessions PowerShell, dans VS Code lancé depuis le menu "
        "Démarrer, et dans Claude Code CLI."
    )
    add_code_block(
        doc,
        '''# Persistant : CA bundle pour Node.js, Bun, Python requests
[Environment]::SetEnvironmentVariable(
    "NODE_EXTRA_CA_CERTS",
    "$env:USERPROFILE\\.mitmproxy\\mitmproxy-ca-cert.pem",
    "User")
[Environment]::SetEnvironmentVariable(
    "SSL_CERT_FILE",
    "$env:USERPROFILE\\.mitmproxy\\mitmproxy-ca-cert.pem",
    "User")

# Cle Anthropic
[Environment]::SetEnvironmentVariable(
    "ANTHROPIC_API_KEY", "sk-ant-api03-...", "User")

# Proxy a activer uniquement pendant les sessions de lab
# A executer dans chaque shell PowerShell de capture :
function llmlab-on  { $env:HTTPS_PROXY = "http://localhost:8888"; $env:HTTP_PROXY = "http://localhost:8888" }
function llmlab-off { Remove-Item Env:HTTPS_PROXY -ErrorAction SilentlyContinue
                       Remove-Item Env:HTTP_PROXY  -ErrorAction SilentlyContinue }

# Pour persister la fonction : l'ajouter au profil PowerShell
notepad $PROFILE''',
        language="powershell",
    )

    doc.add_heading("5.6 Cas particulier : WSL2", level=2)
    doc.add_paragraph(
        "WSL2 fonctionne sur une stack réseau distincte de Windows. "
        "Les variables d'environnement Windows ne sont pas "
        "automatiquement héritées (sauf via WSLENV). La CA "
        "installée dans le Certificate Store Windows ne couvre pas "
        "non plus la stack TLS Linux de WSL."
    )
    add_code_block(
        doc,
        '''# Dans WSL : repeter l'installation Linux complete (voir section 4)
# La CA mitmproxy de WSL est distincte de celle de Windows.

# Pour partager des variables entre Windows et WSL, utiliser WSLENV
# Cote PowerShell, avant de lancer wsl :
$env:WSLENV = "ANTHROPIC_API_KEY:HTTPS_PROXY:NODE_EXTRA_CA_CERTS"

# Dans WSL, le proxy "localhost" Windows est accessible via
# host.docker.internal ou via l'adresse IP de l'interface eth0 du host.
# Methode robuste : recuperer l'IP du host Windows depuis WSL
nameserver=$(grep nameserver /etc/resolv.conf | awk '{print $2}')
export HTTPS_PROXY="http://${nameserver}:8888"''',
        language="bash",
    )

    add_security_box(
        doc,
        "Bun runtime et Claude Code sur Windows",
        [
            "Claude Code v2.x sur Windows embarque le runtime Bun "
            "qui n'utilise pas systématiquement OpenSSL ni le "
            "trust store Windows pour la validation TLS. Plusieurs "
            "issues GitHub documentent des cas où NODE_EXTRA_CA_CERTS "
            "et SSL_CERT_FILE sont ignorés (issues 41157 et 26897).",

            "Depuis Claude Code v2.1.101, le runtime trust par "
            "défaut le store OS via la variable "
            "CLAUDE_CODE_CERT_STORE = bundled,system. Si "
            "l'interception échoue malgré une CA correctement "
            "installée dans Cert:\\LocalMachine\\Root, vérifier "
            "que cette variable n'a pas été surchargée à "
            "bundled seul. En dernier recours, utiliser le mode "
            "reverse avec ANTHROPIC_BASE_URL=http://localhost:8888 "
            "qui contourne le problème de chaîne TLS.",
        ],
    )


def _section_6_smoketest(doc: Document) -> None:
    doc.add_heading("6. Test de fumée HTTP local", level=1)
    doc.add_paragraph(
        "Avant d'attaquer api.anthropic.com en HTTPS, valider la "
        "pipeline avec un test local en HTTP clair contre un "
        "serveur d'écho simple. Cette étape isole les problèmes "
        "potentiels (mitmproxy mal installé, port déjà occupé)."
    )

    add_code_block(
        doc,
        '''# Terminal 1 - lancer mitmproxy en mode reverse devant httpbin.org
mitmdump --mode reverse:https://httpbin.org -p 8888

# Terminal 2 - emettre une requete
curl -sS http://localhost:8888/get
# La reponse JSON est affichee, et mitmproxy console montre la
# requete capturee avec son URL, ses headers, son body.''',
        language="bash",
    )

    doc.add_paragraph(
        "Si le test fonctionne, la pipeline est saine. Si curl "
        "renvoie une erreur de connexion, vérifier que le port 8888 "
        "n'est pas occupé par un autre processus (lsof -i :8888 sur "
        "Linux, Get-NetTCPConnection -LocalPort 8888 sur Windows)."
    )


def _section_7_https_test(doc: Document) -> None:
    doc.add_heading("7. Test HTTPS sur api.anthropic.com", level=1)
    doc.add_paragraph(
        "Cette étape valide que la CA mitmproxy est correctement "
        "installée dans le trust store et que le déchiffrement TLS "
        "fonctionne sur un endpoint réel."
    )

    add_code_block(
        doc,
        '''# Terminal 1 - lancer mitmproxy en mode regular avec l'addon
mitmdump -p 8888 \\
    -s llm_traffic_logger.py \\
    --set llm_log_path=./traffic.jsonl

# Terminal 2 - emettre une requete Anthropic via le proxy
export HTTPS_PROXY=http://localhost:8888
export ANTHROPIC_API_KEY=sk-ant-...

curl -sS https://api.anthropic.com/v1/messages \\
  -H "x-api-key: $ANTHROPIC_API_KEY" \\
  -H "anthropic-version: 2023-06-01" \\
  -H "content-type: application/json" \\
  -d '{
    "model": "claude-opus-4-7",
    "max_tokens": 64,
    "messages": [
      {"role": "user", "content": "Dis bonjour en une phrase."}
    ]
  }'

# Sur Windows PowerShell, utiliser curl.exe (et non l'alias curl)
# et echapper les guillemets selon le shell.''',
        language="bash",
    )

    doc.add_paragraph(
        "La réponse JSON doit contenir un bloc usage avec "
        "input_tokens et output_tokens, et une entrée doit "
        "apparaître dans traffic.jsonl. Si la requête échoue avec "
        "une erreur TLS (SELF_SIGNED_CERT_IN_CHAIN, "
        "CERT_UNTRUSTED), la CA mitmproxy n'est pas reconnue par "
        "curl ; refaire l'étape 4.3 (Linux) ou 5.4 (Windows)."
    )

    add_tip_box(
        doc,
        "test sans clé Anthropic",
        [
            "Pour valider la pipeline TLS sans consommer de "
            "tokens, on peut tester avec une clé invalide :",
            "  curl -sS https://api.anthropic.com/v1/messages \\\n"
            "    -H \"x-api-key: sk-ant-invalid\" \\\n"
            "    -H \"anthropic-version: 2023-06-01\" \\\n"
            "    -H \"content-type: application/json\" \\\n"
            "    -d '{\"model\":\"claude-opus-4-7\",\"max_tokens\":1,\"messages\":[]}'",
            "L'API renvoie une erreur 401 authentication_error qui "
            "prouve que la connexion TLS a abouti et que le "
            "déchiffrement par mitmproxy fonctionne.",
        ],
    )


def _section_8_clients(doc: Document) -> None:
    doc.add_heading("8. Configuration des clients", level=1)

    doc.add_heading("8.1 Claude Code CLI", level=2)
    doc.add_paragraph(
        "Claude Code CLI respecte HTTPS_PROXY et NODE_EXTRA_CA_CERTS "
        "depuis sa version 1.x. Depuis la v2.1.101, le trust store "
        "OS est consulté par défaut via CLAUDE_CODE_CERT_STORE = "
        "bundled,system, ce qui rend l'interception transparente "
        "une fois la CA installée."
    )
    add_code_block(
        doc,
        '''# Mode regular (recommande)
export HTTPS_PROXY=http://localhost:8888
export ANTHROPIC_API_KEY=sk-ant-...
claude "Refactore le fichier auth.py pour utiliser argon2"

# Mode reverse alternatif (sans manipuler le trust store)
export ANTHROPIC_BASE_URL=http://localhost:8888
export ANTHROPIC_API_KEY=sk-ant-...
claude "Refactore le fichier auth.py pour utiliser argon2"

# En cas de probleme persistant avec le runtime Bun sous Windows,
# forcer la consultation du store OS uniquement :
export CLAUDE_CODE_CERT_STORE=system''',
        language="bash",
    )

    doc.add_heading("8.2 Continue (extension VS Code)", level=2)
    doc.add_paragraph(
        "Continue lit sa configuration depuis ~/.continue/config.json "
        "ou config.yaml. Deux approches pour le proxy."
    )
    add_code_block(
        doc,
        '''// ~/.continue/config.json - mode reverse
{
  "models": [
    {
      "title": "Claude Opus 4.7 via proxy",
      "provider": "anthropic",
      "model": "claude-opus-4-7",
      "apiBase": "http://localhost:8888",
      "apiKey": "sk-ant-..."
    }
  ]
}''',
        language="json",
    )
    doc.add_paragraph(
        "Pour le mode regular, ne pas définir apiBase et exporter "
        "HTTPS_PROXY avant de lancer VS Code. Si VS Code est "
        "démarré depuis le menu Démarrer Windows sans terminal "
        "préalable, utiliser le profil PowerShell ou les variables "
        "User scope pour persister HTTPS_PROXY."
    )

    doc.add_heading("8.3 Cline (extension VS Code)", level=2)
    doc.add_paragraph(
        "Cline (anciennement Claude Dev) accepte la configuration "
        "Anthropic dans son panneau de paramètres. Le champ "
        "Anthropic Base URL permet l'override pour le mode reverse. "
        "Pour le mode regular, exporter HTTPS_PROXY puis lancer "
        "code depuis ce terminal."
    )

    doc.add_heading("8.4 SDK Python officiel", level=2)
    add_code_block(
        doc,
        '''# Mode regular - sans modification du code
# Suffit d'exporter HTTPS_PROXY avant de lancer le script
export HTTPS_PROXY=http://localhost:8888

python3 - <<'PYEOF2'
import os
from anthropic import Anthropic
client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
msg = client.messages.create(
    model="claude-opus-4-7",
    max_tokens=128,
    messages=[{"role": "user", "content": "Ping"}],
)
print(msg.content[0].text)
PYEOF2

# Mode reverse - override base_url dans le code
python3 - <<'PYEOF2'
import os
from anthropic import Anthropic
client = Anthropic(
    api_key=os.environ["ANTHROPIC_API_KEY"],
    base_url="http://localhost:8888",
)
msg = client.messages.create(
    model="claude-opus-4-7",
    max_tokens=128,
    messages=[{"role": "user", "content": "Ping"}],
)
print(msg.content[0].text)
PYEOF2''',
        language="bash",
    )


def _section_9_addon(doc: Document) -> None:
    doc.add_heading("9. Lancement de l'addon de capture", level=1)
    doc.add_paragraph(
        "L'addon llm_traffic_logger.py est fourni dans la note "
        "technique principale (Annexe A de "
        "llm_stateless_architecture.docx). Le copier dans le "
        "répertoire de travail."
    )

    add_code_block(
        doc,
        '''# Recuperation de l'addon (ajuster le chemin source)
cp /path/to/llm_traffic_logger.py ./

# Lancement en mode regular
mitmdump -p 8888 \\
    -s llm_traffic_logger.py \\
    --set llm_log_path=./traffic.jsonl

# Variantes utiles
# Avec mitmweb (interface graphique sur localhost:8081)
mitmweb -p 8888 \\
    -s llm_traffic_logger.py \\
    --set llm_log_path=./traffic.jsonl

# Avec verbosity accrue pour debug
mitmdump -p 8888 \\
    -s llm_traffic_logger.py \\
    --set llm_log_path=./traffic.jsonl \\
    --set termlog_verbosity=debug''',
        language="bash",
    )

    doc.add_paragraph(
        "Vérifier qu'une première entrée arrive dans traffic.jsonl "
        "dès l'émission d'une requête depuis le client. Le format "
        "JSONL est ingestible directement dans jq, pandas, Splunk "
        "ou OpenSearch."
    )

    add_code_block(
        doc,
        '''# Inspection rapide de la derniere entree avec jq
tail -n 1 traffic.jsonl | jq '{
  path: .path,
  status: .response_status,
  duration: (.ts_response - .ts_request),
  input_tokens: (.sse_events // []) |
                map(select(.data.type == "message_start")) |
                map(.data.message.usage.input_tokens) |
                max,
  events: (.sse_events // []) | length
}'

# Comptage rapide des tours captures
wc -l traffic.jsonl''',
        language="bash",
    )


def _section_10_scenarios(doc: Document, diagrams: dict) -> None:
    doc.add_heading("10. Scénarios d'observation", level=1)
    doc.add_paragraph(
        "Quatre scénarios à dérouler dans l'ordre. Chaque scénario "
        "produit une signature distinctive dans le JSONL qui "
        "valide à la fois l'interception et la compréhension du "
        "mécanisme observé."
    )
    add_figure(doc, *diagrams["lab04_scenarios"])

    doc.add_heading("10.1 Scénario 1 - Premier message", level=2)
    doc.add_paragraph(
        "Validation du chemin nominal sans cache_control. Une "
        "seule requête, un seul tour. Permet d'observer la "
        "structure complète du flux SSE."
    )
    add_code_block(
        doc,
        '''curl -sS https://api.anthropic.com/v1/messages \\
  -H "x-api-key: $ANTHROPIC_API_KEY" \\
  -H "anthropic-version: 2023-06-01" \\
  -H "content-type: application/json" \\
  -d '{
    "model": "claude-opus-4-7",
    "max_tokens": 256,
    "stream": true,
    "messages": [
      {"role": "user",
       "content": "Decris en trois lignes ce qu'est OWASP LLM Top 10."}
    ]
  }'

# Verification de la trace
tail -n 1 traffic.jsonl | jq '.sse_events | length'
# Doit etre superieur a 5 (message_start, content_block_start,
# plusieurs content_block_delta, content_block_stop, message_delta,
# message_stop)''',
        language="bash",
    )

    doc.add_heading("10.2 Scénario 2 - Cache write puis cache hit", level=2)
