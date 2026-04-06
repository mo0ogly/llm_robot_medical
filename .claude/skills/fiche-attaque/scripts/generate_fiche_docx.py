"""
Generate a complete 10-section Fiche d'Attaque .docx from structured markdown sections.

Usage:
    python generate_fiche_docx.py --metadata metadata.json --sections sections.json --output fiche.docx

This script is called by the fiche-attaque skill after the 3 agents produce their sections.
It assembles everything into a standardized .docx document.
"""

import argparse
import io
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# --- Unicode delta symbols ---
DELTA = {
    "delta0": "\u03b4\u2070",
    "delta1": "\u03b4\u00b9",
    "delta2": "\u03b4\u00b2",
    "delta3": "\u03b4\u00b3",
}


def normalize_delta(text: str) -> str:
    """Replace ASCII delta references with Unicode symbols."""
    for ascii_form, unicode_form in DELTA.items():
        text = text.replace(ascii_form, unicode_form)
        text = text.replace("delta-0", DELTA["delta0"])
        text = text.replace("delta-1", DELTA["delta1"])
        text = text.replace("delta-2", DELTA["delta2"])
        text = text.replace("delta-3", DELTA["delta3"])
    return text


def add_metadata_table(doc, metadata: dict):
    """Add the metadata header table."""
    fields = [
        ("Identifiant", metadata.get("id", "")),
        ("Categorie", metadata.get("category", "")),
        ("Couche ciblee", normalize_delta(metadata.get("target_delta", "delta1"))),
        ("Conjecture", metadata.get("conjecture", "-") or "-"),
        ("Chain ID", metadata.get("chain_id", "-") or "-"),
        ("MITRE ATT&CK", metadata.get("mitre", "T1059")),
        ("OWASP LLM", metadata.get("owasp", "LLM01:2025 - Prompt Injection")),
        ("Ref. these", metadata.get("thesis_ref", "")),
    ]
    table = doc.add_table(rows=len(fields), cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(fields):
        cell_label = table.rows[i].cells[0]
        cell_value = table.rows[i].cells[1]
        cell_label.text = label
        cell_value.text = str(value) if value else "-"
        # Bold labels
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def add_markdown_section(doc, markdown_text: str, base_level: int = 1):
    """Convert simple markdown to docx paragraphs.

    Handles: # headings, - bullets, **bold**, regular text, tables (|).
    """
    if not markdown_text:
        return

    lines = markdown_text.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Table detection
        if stripped.startswith("|") and "|" in stripped[1:]:
            # Skip separator rows
            if all(c in "|-: " for c in stripped):
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        elif in_table:
            # Flush table
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Headings
        if stripped.startswith("#### "):
            doc.add_heading(normalize_delta(stripped[5:]), level=min(base_level + 3, 4))
        elif stripped.startswith("### "):
            doc.add_heading(normalize_delta(stripped[4:]), level=min(base_level + 2, 4))
        elif stripped.startswith("## "):
            doc.add_heading(normalize_delta(stripped[3:]), level=min(base_level + 1, 3))
        elif stripped.startswith("# "):
            doc.add_heading(normalize_delta(stripped[2:]), level=base_level)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(normalize_delta(stripped[2:]), style="List Bullet")
        elif stripped.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            doc.add_paragraph(normalize_delta(stripped), style="List Number")
        elif stripped:
            p = doc.add_paragraph()
            _add_formatted_text(p, normalize_delta(stripped))
        # Skip empty lines

    if in_table:
        _flush_table(doc, table_rows)


def _flush_table(doc, rows):
    """Create a docx table from collected rows."""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols, style="Table Grid")
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < max_cols:
                table.rows[i].cells[j].text = normalize_delta(cell_text)
    # Bold first row (header)
    if rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def _add_formatted_text(paragraph, text: str):
    """Parse **bold** markers in text."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            if i % 2 == 1:  # Odd indices are bold
                run.bold = True


def generate_fiche(
    metadata: dict,
    sections: dict,
    output_path: str = None,
) -> bytes:
    """Generate a complete 10-section Fiche d'Attaque.

    Args:
        metadata: dict with id, name, category, target_delta, conjecture, chain_id, etc.
        sections: dict with keys 'section_1' through 'section_10', 'annexe_a', 'annexe_b'
                  Each value is markdown text.
        output_path: optional file path to save the .docx

    Returns:
        bytes of the .docx file
    """
    doc = Document()

    # --- Style setup ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading("FICHE D'ATTAQUE", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Template name ---
    name = metadata.get("name", "Unknown Template")
    subtitle = doc.add_heading(name, level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- One-line description ---
    desc = metadata.get("description", "")
    if desc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(desc)
        run.italic = True

    doc.add_paragraph()  # Spacer

    # --- Metadata table ---
    add_metadata_table(doc, metadata)
    doc.add_paragraph()  # Spacer

    # --- 10 Sections ---
    section_titles = {
        "section_1": "1. Contexte et Modele de Menace",
        "section_2": "2. Prompt d'Injection",
        "section_3": "3. Cadre Mathematique",
        "section_4": "4. Analyse IA - Mecanismes de Refus",
        "section_5": "5. Score SVC - Evaluation Multidimensionnelle",
        "section_6": "6. Parametres Testes et Resultats",
        "section_7": "7. Axes de Defense et Contre-mesures",
        "section_8": "8. Analyse Corrective - Ce Qui Aurait Du Etre Fait",
        "section_9": "9. Etat de l'Art et References",
        "section_10": "10. Verdict et Positionnement dans la These",
        "section_11": "11. Analyse Scientifique et Axes de Recherche",
    }

    for key, title_text in section_titles.items():
        doc.add_heading(normalize_delta(title_text), level=1)
        content = sections.get(key, "")
        if content:
            add_markdown_section(doc, content, base_level=2)
        else:
            doc.add_paragraph("[Section a completer]")
        doc.add_paragraph()  # Spacer between sections

    # --- Annexes ---
    doc.add_heading("ANNEXE A - Guide de Remplissage du Template", level=1)
    annexe_a = sections.get("annexe_a", "")
    if annexe_a:
        add_markdown_section(doc, annexe_a, base_level=2)
    else:
        doc.add_paragraph("[Annexe a completer]")

    doc.add_heading("ANNEXE B - Glossaire Mathematique", level=1)
    annexe_b = sections.get("annexe_b", "")
    if annexe_b:
        add_markdown_section(doc, annexe_b, base_level=2)
    else:
        _add_default_glossary(doc)

    # --- Save ---
    buffer = io.BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()

    if output_path:
        Path(output_path).write_bytes(content)
        print("OK: " + output_path)

    return content


def _add_default_glossary(doc):
    """Add the standard mathematical glossary as Annexe B."""
    glossary = [
        ("\u03b4\u2070", "Couche de base - alignement RLHF implicite"),
        ("\u03b4\u00b9", "Couche 1 - defense par system prompt"),
        ("\u03b4\u00b2", "Couche 2 - sanitization des entrees (RAG Shield)"),
        ("\u03b4\u00b3", "Couche 3 - verification multi-agent (LLM Judge)"),
        ("Sep(M)", "Score de separation instruction/donnee (Zverev et al., 2025)"),
        ("Reachable(S, p)", "Fonction d'atteignabilite - True si p peut violer Pi dans S"),
        ("P(detect)", "Probabilite de detection d'un pattern d'injection"),
        ("ASR", "Attack Success Rate - taux de succes sur N essais"),
        ("SVC", "Score de Vulnerability Composite (6 dimensions)"),
        ("DY-AGENT", "Modele Dolev-Yao adapte aux systemes agentiques"),
        ("Pi", "Ensemble des politiques de securite du systeme S"),
        ("N >= 30", "Seuil de validite statistique pour Sep(M)"),
    ]
    table = doc.add_table(rows=len(glossary) + 1, cols=2, style="Table Grid")
    table.rows[0].cells[0].text = "Symbole / Notation"
    table.rows[0].cells[1].text = "Definition"
    for p in table.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for p in table.rows[0].cells[1].paragraphs:
        for r in p.runs:
            r.bold = True
    for i, (symbol, definition) in enumerate(glossary):
        table.rows[i + 1].cells[0].text = symbol
        table.rows[i + 1].cells[1].text = definition


# --- CLI entry point ---
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Fiche d'Attaque .docx")
    parser.add_argument("--metadata", required=True, help="Path to metadata JSON")
    parser.add_argument("--sections", required=True, help="Path to sections JSON")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    with open(args.metadata) as f:
        meta = json.load(f)
    with open(args.sections) as f:
        secs = json.load(f)

    generate_fiche(meta, secs, args.output)
