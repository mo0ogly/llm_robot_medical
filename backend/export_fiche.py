"""
Export attack sheet (fiche d'attaque) as .docx document.
Generates a structured thesis-grade document from template data.
"""

import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_fiche(
    template_data: dict,
    versions: list = None,
    help_content: str = "",
    detection_baseline: dict = None,
) -> io.BytesIO:
    """Generate a .docx attack sheet from template data.

    Args:
        template_data: dict with id, name, category, template, variables, detection_profile
        versions: list of version dicts (for section 8 corrective analysis)
        help_content: markdown help content (for section details)
        detection_baseline: dict of baseline detection probabilities

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # ------- Title -------
    title = doc.add_heading("FICHE D'ATTAQUE", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Template name
    doc.add_heading(template_data.get("name", "Unknown"), level=1)

    # ------- Metadata table -------
    table = doc.add_table(rows=6, cols=2, style="Table Grid")
    metadata = [
        ("Identifiant", template_data.get("id", "")),
        ("Categorie", template_data.get("category", "")),
        ("Couche ciblee", template_data.get("target_delta", "delta1")),
        ("Conjecture", template_data.get("conjecture", "-")),
        ("Chain ID", template_data.get("chain_id", "-")),
        ("OWASP LLM", "LLM01:2025 --- Prompt Injection"),
    ]
    for i, (label, value) in enumerate(metadata):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value) if value else "-"

    # ------- Section 2: Prompt template -------
    doc.add_heading("2. Prompt d'Injection", level=1)
    doc.add_heading("2.1 Template brut", level=2)
    doc.add_paragraph(template_data.get("template", ""), style="No Spacing")

    # Variables table
    variables = template_data.get("variables", {})
    if variables:
        doc.add_heading("2.2 Variables d'instanciation", level=2)
        var_table = doc.add_table(
            rows=len(variables) + 1, cols=2, style="Table Grid"
        )
        var_table.rows[0].cells[0].text = "Variable"
        var_table.rows[0].cells[1].text = "Valeur"
        for i, (k, v) in enumerate(variables.items()):
            var_table.rows[i + 1].cells[0].text = "{{" + k + "}}"
            var_table.rows[i + 1].cells[1].text = str(v)

    # ------- Section 5: SVC Score -------
    doc.add_heading("5. Score SVC --- Evaluation Multidimensionnelle", level=1)
    detection_profile = template_data.get("detection_profile", {})
    patterns = [
        "caps",
        "negation",
        "token_fictif",
        "tool_direct",
        "coercion",
        "xml_fictif",
    ]
    labels = {
        "caps": "CAPS / Urgency",
        "negation": "Explicit Negation",
        "token_fictif": "Fictitious Token",
        "tool_direct": "Direct Tool Invoke",
        "coercion": "Coercive Framing",
        "xml_fictif": "Fictitious XML/HL7",
    }
    if detection_profile:
        svc_table = doc.add_table(rows=7, cols=3, style="Table Grid")
        svc_table.rows[0].cells[0].text = "Pattern"
        svc_table.rows[0].cells[1].text = "P(detect)"
        svc_table.rows[0].cells[2].text = "Reference"
        for i, p in enumerate(patterns):
            svc_table.rows[i + 1].cells[0].text = labels.get(p, p)
            svc_table.rows[i + 1].cells[1].text = str(
                detection_profile.get(p, "-")
            )
            svc_table.rows[i + 1].cells[2].text = ""

    # ------- Section 8: Corrective Analysis (if versions exist) -------
    if versions:
        doc.add_heading("8. Analyse Corrective --- Versions", level=1)
        for idx, ver in enumerate(versions):
            doc.add_heading(
                ver.get("version_label", "V" + str(idx + 2)), level=2
            )
            doc.add_paragraph(
                ver.get("template", ""), style="No Spacing"
            )

            # Detection comparison
            ver_profile = ver.get("detection_profile", {})
            if ver_profile and detection_baseline:
                doc.add_heading("Comparaison de detection", level=3)
                comp_table = doc.add_table(
                    rows=len(patterns) + 2, cols=4, style="Table Grid"
                )
                comp_table.rows[0].cells[0].text = "Pattern"
                comp_table.rows[0].cells[1].text = "P(detect) Baseline"
                comp_table.rows[0].cells[2].text = "P(detect) Evolved"
                comp_table.rows[0].cells[3].text = "Delta"
                product_base = 1.0
                product_evol = 1.0
                for i, p in enumerate(patterns):
                    base_p = (
                        detection_baseline.get("patterns", {})
                        .get(p, {})
                        .get("p_detect", 0)
                    )
                    evol_p = ver_profile.get(p, base_p)
                    delta = evol_p - base_p
                    comp_table.rows[i + 1].cells[0].text = labels.get(p, p)
                    comp_table.rows[i + 1].cells[1].text = (
                        str(round(base_p * 100)) + "%"
                    )
                    comp_table.rows[i + 1].cells[2].text = (
                        str(round(evol_p * 100)) + "%"
                    )
                    comp_table.rows[i + 1].cells[3].text = (
                        ("+" if delta >= 0 else "")
                        + str(round(delta * 100))
                        + "%"
                    )
                    product_base *= 1 - base_p
                    product_evol *= 1 - evol_p
                # Cumulative row
                last = len(patterns) + 1
                comp_table.rows[last].cells[0].text = "P(detect) cumul."
                comp_table.rows[last].cells[1].text = (
                    str(round((1 - product_base) * 100, 4)) + "%"
                )
                comp_table.rows[last].cells[2].text = (
                    str(round((1 - product_evol) * 100, 2)) + "%"
                )
                comp_table.rows[last].cells[3].text = "-"

    # ------- Section 10: Verdict -------
    doc.add_heading("10. Verdict", level=1)
    doc.add_paragraph("[A REMPLIR apres execution des campagnes]")

    # ------- Annexe: Help content -------
    if help_content:
        doc.add_heading("ANNEXE : Documentation technique", level=1)
        for line in help_content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
