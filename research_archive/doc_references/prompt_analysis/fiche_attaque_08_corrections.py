#!/usr/bin/env python3
"""Apply anti-confabulation corrections to fiche #08 (Extortion Tool Hijack).

Round 1 (numerical) and round 2 (attribution/date/claim) corrections surfaced by
the anti-confabulation audit and the scoped reference verification:

Round 1 -- numerical:
    1. "6 ordres de grandeur" -> "Environ 2.5 ordres" (verified by recompute).
    2. CI: [0%, 9.5%] is the one-sided 95% bound; the two-sided 95% Clopper-Pearson
       bound for 0/30 is 11.6%. Projected status of the result is made explicit.
    3. Sep(M)=1.0 over N=30 was shown as a measurement while section 6 is [REMPLIR];
       relabelled as a projection.

Round 2 -- attribution / date / claim (scoped verification, 2026-05-21):
    4. "Wei et al." -> "Qi et al." : "Safety Alignment ... Tokens Deep" is Qi et al.
       (ICLR 2025, arXiv:2406.05946 = corpus P018), not Wei et al.
    5. The fine-tuning fragility result ("100 exemples / GSM8K") is Qi et al. 2023
       (arXiv:2310.03693), not NDSS 2025.
    6. Schulhoff et al. corrected to 2023 (EMNLP), not 2024.
    7. The "aucune taxonomie" claim is bounded to the consulted taxonomies (HUMILITY GATE).

The patch is run-aware (targets sit in single runs, so formatting is preserved) and
idempotent (re-running is a no-op once corrections are in). Note text deliberately
avoids the literal "Wei et al." token so the blanket replacement cannot clobber it
on a re-run.

Usage:
    python fiche_attaque_08_corrections.py \
        --input  fiche_attaque_08_extortion_tool_hijack.docx \
        --output fiche_attaque_08_extortion_tool_hijack.docx
"""

from __future__ import annotations

import argparse
import logging
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
logger = logging.getLogger("fiche08_corrections")

# Ordered (old, new) substring replacements applied at run level.
REPLACEMENTS: list[tuple[str, str]] = [
    # Round 1 -- numerical
    (
        "6 ordres de grandeur au-dessus du #01",
        "Environ 2.5 ordres de grandeur au-dessus du #01",
    ),
    (
        "Clopper-Pearson pour ASR = 0/30",
        "Clopper-Pearson pour un resultat projete ASR = 0/30",
    ),
    (
        "[0%, 9.5%]",
        "[0%, 11.6%] (bilateral ; borne unilaterale 95% = 9.5%)",
    ),
    (
        "Sep(M) = 1.0 sur N = 30 essais",
        "Sep(M) = 1.0 attendu sur N = 30 essais "
        "[PROJECTION : campagne non encore executee, cf. section 6]",
    ),
    # Round 2 -- attribution / date / claim
    (
        "Wei et al.",
        "Qi et al.",
    ),
    (
        "Schulhoff et al. (2024)",
        "Schulhoff et al. (2023)",
    ),
    (
        "Fine-tuning sur 100 exemples malveillants",
        "Fine-tuning adversarial (Qi et al. 2023, environ 10 exemples adversariaux)",
    ),
    (
        "fragile au fine-tuning (NDSS 2025)",
        "fragile au fine-tuning (Qi et al. 2023, arXiv:2310.03693 ; NDSS 2025)",
    ),
    (
        "aucune taxonomie comme technique viable",
        "aucune des taxonomies consultees (Schulhoff 2023, CrowdStrike 2026) "
        "comme technique viable",
    ),
]

# Revision notes appended at the end of the document (idempotent, one each).
NOTES: list[tuple[str, str]] = [
    (
        "Note de revision (2026-05-21)",
        "Note de revision (2026-05-21) : corrections anti-confabulation appliquees "
        "(ecart d'ordres de grandeur, IC95 bilateral, statut projete de la section 3.5). "
        "Detail : research_notes/AEGIS-AUDIT-FICHE-08_anti-confabulation.md",
    ),
    (
        "Note de revision 2 (2026-05-21)",
        "Note de revision 2 (2026-05-21) : corrections d'attribution anti-confabulation. "
        "'Safety Alignment Should Be Made More Than Just a Few Tokens Deep' est de Qi et al. "
        "(ICLR 2025, arXiv:2406.05946 = corpus P018), attribue par erreur a Wei dans la "
        "version initiale ; le resultat fine-tuning (100 exemples / GSM8K) est de Qi et al. "
        "2023 (arXiv:2310.03693), non NDSS 2025 ; Schulhoff et al. corrige en 2023 (EMNLP) ; "
        "la claim 'aucune taxonomie' est bornee aux taxonomies consultees. "
        "Detail : research_notes/AEGIS-SCOPED-VERIF_fiche08-refs_2026-05-21.md",
    ),
]


def iter_all_paragraphs(parent) -> Iterator[Paragraph]:
    """Yield every paragraph in the document body, including nested table cells."""
    yield from parent.paragraphs
    tables: list[Table] = list(parent.tables)
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)


def apply_replacements(document: Document) -> dict[str, int]:
    """Apply run-level replacements; fall back to cross-run for split matches."""
    counts = {old: 0 for old, _ in REPLACEMENTS}
    for paragraph in iter_all_paragraphs(document):
        runs = paragraph.runs
        for old, new in REPLACEMENTS:
            # Fast path: substring contained within a single run.
            for run in runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    counts[old] += 1
            # Slow path: substring split across runs (collapse to first run).
            joined = "".join(r.text for r in runs)
            if old in joined and runs:
                runs[0].text = joined.replace(old, new)
                for r in runs[1:]:
                    r.text = ""
                counts[old] += 1
                runs = paragraph.runs
    return counts


def ensure_note(document: Document, marker: str, text: str) -> bool:
    """Append a note paragraph once (idempotent); return True if it was added."""
    if any(marker in p.text for p in document.paragraphs):
        return False
    document.add_paragraph(text)
    return True


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    here = Path(__file__).resolve().parent
    default_doc = here / "fiche_attaque_08_extortion_tool_hijack.docx"
    parser.add_argument("--input", type=Path, default=default_doc)
    parser.add_argument("--output", type=Path, default=default_doc)
    args = parser.parse_args()

    logger.info("Loading %s", args.input)
    document = Document(str(args.input))

    counts = apply_replacements(document)
    for old, hits in counts.items():
        level = logging.INFO if hits else logging.WARNING
        logger.log(level, "[%d hit] %s", hits, old[:48])

    for marker, text in NOTES:
        added = ensure_note(document, marker, text)
        logger.info("Note '%s' added: %s", marker, added)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(args.output))
    logger.info("Saved corrected document to %s", args.output)


if __name__ == "__main__":
    main()
