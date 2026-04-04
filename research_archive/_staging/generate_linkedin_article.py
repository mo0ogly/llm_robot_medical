#!/usr/bin/env python3
"""
Generate LinkedIn article: Multi-Agent Research Team for Doctoral Thesis
Format: Word .docx with integrated workflow schemas
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_border_to_paragraph(paragraph, color="4472C4"):
    """Add a left border to a paragraph (for callout boxes)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def create_article():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ========================================================================
    # TITLE
    # ========================================================================
    title = doc.add_heading('', level=0)
    run = title.add_run("9 Agents IA Autonomes pour une These Doctorale :\nComment j'ai Monte une Equipe de Recherche qui Communique, Itere et Apprend")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("De la recherche bibliographique manuelle a un swarm d'agents specialises\nqui decouvrent, analysent, enseignent et preparent 46 articles en 45 minutes")
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Author line
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Florent Pizzini | Doctorant ENS 2026 | Securite IA & Prompt Injection")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.bold = True

    doc.add_paragraph("")  # spacer

    # ========================================================================
    # HOOK / INTRODUCTION
    # ========================================================================
    doc.add_heading("Le Probleme : Une These, 46 Articles, Zero Temps", level=1)

    p = doc.add_paragraph()
    p.add_run("Imaginez : ").bold = True
    p.add_run(
        "vous preparez une these doctorale sur la securite des LLM en contexte medical. "
        "Votre directeur de these attend une couverture bibliographique exhaustive — "
        "pas 10 articles, pas 20, mais une cartographie complete du domaine prompt injection "
        "de 2023 a 2026. Chaque article doit etre resume en francais, ses formules "
        "mathematiques expliquees pour un public bac+2, ses menaces cartographiees, "
        "ses techniques d'attaque reproductibles, et le tout organise dans un systeme "
        "interrogeable par RAG."
    )

    p = doc.add_paragraph()
    p.add_run(
        "Manuellement ? Comptez 2-3 heures par article. Pour 46 articles : "
    )
    run = p.add_run("~120 heures de travail repetitif.")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p = doc.add_paragraph()
    p.add_run("Ma solution : ").bold = True
    p.add_run(
        "un swarm de 9 agents IA autonomes, chacun specialise dans un role de recherche, "
        "qui communiquent entre eux via un systeme de memoire partagee, iterent a chaque "
        "execution, et produisent un travail de qualite doctorale — en ")
    run = p.add_run("45 minutes.")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    # ========================================================================
    # SECTION 1: THE CONCEPT
    # ========================================================================
    doc.add_heading("1. Le Concept : Un Laboratoire de Recherche Virtuel", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "L'idee centrale est simple mais puissante : plutot que d'utiliser un seul agent IA "
        "generique pour tout faire, j'ai decompose le travail de recherche en "
    )
    p.add_run("9 roles specialises").bold = True
    p.add_run(
        ", exactement comme dans un laboratoire de recherche reel. "
        "Chaque agent a son expertise, ses outils, ses criteres de succes, "
        "et surtout — ils communiquent entre eux."
    )

    # Agent roles table
    doc.add_heading("Les 9 Roles du Laboratoire", level=2)

    table = doc.add_table(rows=10, cols=4)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Agent', 'Role', 'Analogie Labo', 'Livrable Principal']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    agents_data = [
        ["COLLECTOR", "Decouverte d'articles", "Documentaliste", "46 articles + metadonnees"],
        ["ANALYST", "Resume FR + gaps", "Chercheur principal", "34 analyses de 500 mots"],
        ["MATHEUX", "Glossaire formules", "Mathematicien", "22 formules expliquees + DAG"],
        ["CYBERSEC", "Modeles de menaces", "Analyste securite", "Matrice MITRE ATT&CK"],
        ["WHITEHACKER", "Techniques d'attaque", "Pentester", "18 techniques + 12 PoC"],
        ["LIBRARIAN", "Organisation fichiers", "Bibliothecaire", "MANIFEST + INDEX"],
        ["MATHTEACHER", "Cours de maths FR", "Professeur", "7 modules + 34 exercices"],
        ["SCIENTIST", "Axes de recherche", "Directeur de these", "8 axes + 6 conjectures"],
        ["CHUNKER", "Preparation RAG", "Ingenieur donnees", "290 chunks ChromaDB"],
    ]

    for i, row_data in enumerate(agents_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_text
            for p in table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph("")

    # ========================================================================
    # SECTION 2: WORKFLOW SCHEMA
    # ========================================================================
    doc.add_heading("2. Architecture du Workflow Multi-Agent", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Le secret n'est pas d'avoir 9 agents — c'est de les orchestrer intelligemment. "
        "Le workflow suit 5 phases avec des portes de validation (gates) qui garantissent "
        "que chaque agent recoit des donnees fiables de ses predecesseurs."
    )

    # WORKFLOW DIAGRAM as formatted text
    doc.add_heading("Schema du Pipeline Orchestrateur", level=2)

    # Phase diagram as a table
    workflow_table = doc.add_table(rows=6, cols=3)
    workflow_table.style = 'Light Grid Accent 1'
    workflow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    wf_headers = ['Phase', 'Agents', 'Gate de Sortie']
    for i, h in enumerate(wf_headers):
        workflow_table.rows[0].cells[i].text = h
        for p in workflow_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    wf_data = [
        ["P1 - Decouverte", "COLLECTOR\n(6 requetes paralleles)", ">=20 articles uniques, zero doublons"],
        ["P2 - Analyse\n(parallele x4)", "ANALYST + MATHEUX\n+ CYBERSEC + WHITEHACKER", "Toutes analyses completees\n+ rapports valides"],
        ["P3 - Organisation", "LIBRARIAN\n(fichiers + index)", "MANIFEST complet,\nzero orphelins"],
        ["P4 - Synthese\n(parallele x2)", "MATHTEACHER\n+ SCIENTIST", "Curriculum + axes\nde recherche valides"],
        ["P5 - Ingestion", "CHUNKER\n(RAG prep)", "290 chunks, script\nd'ingestion teste"],
    ]

    for i, row_data in enumerate(wf_data):
        for j, cell_text in enumerate(row_data):
            workflow_table.rows[i+1].cells[j].text = cell_text
            for p in workflow_table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph("")

    # Inter-agent communication diagram
    doc.add_heading("Schema des Communications Inter-Agents", level=2)

    comm_diagram = doc.add_paragraph()
    comm_diagram.paragraph_format.space_before = Pt(6)
    comm_diagram.paragraph_format.space_after = Pt(6)

    diagram_text = (
        "                    ORCHESTRATEUR\n"
        "                         |\n"
        "              +----------+----------+\n"
        "              |                     |\n"
        "         COLLECTOR              MEMORY_STATE.md\n"
        "         (WebSearch)            (Memoire partagee)\n"
        "              |                     ^\n"
        "              v                     |\n"
        "    papers_phase1.json    <--- Lecture/Ecriture par TOUS --->\n"
        "    papers_phase2.json\n"
        "              |\n"
        "    +---------+---------+---------+\n"
        "    |         |         |         |\n"
        " ANALYST   MATHEUX  CYBERSEC  WHITEHACKER\n"
        " (Resume)  (Maths)  (Menaces) (Exploits)\n"
        "    |         |         |         |\n"
        "    +---------+---------+---------+\n"
        "              |\n"
        "          LIBRARIAN\n"
        "    (Indexes + Filesystem)\n"
        "              |\n"
        "    +---------+---------+\n"
        "    |                   |\n"
        " MATHTEACHER        SCIENTIST\n"
        " (Curriculum)    (Axes recherche)\n"
        "    |                   |\n"
        "    +---------+---------+\n"
        "              |\n"
        "           CHUNKER\n"
        "     (RAG ChromaDB prep)\n"
    )

    run = comm_diagram.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_paragraph("")

    # ========================================================================
    # SECTION 3: INTER-AGENT COMMUNICATION
    # ========================================================================
    doc.add_heading("3. Communication Inter-Agents : Le Coeur du Systeme", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "La veritable innovation n'est pas l'utilisation d'agents IA — c'est la facon "
        "dont ils communiquent et se coordonnent. Trois mecanismes cles :"
    )

    # 3.1 Shared Memory
    doc.add_heading("3.1 Memoire Partagee (MEMORY_STATE.md)", level=2)

    p = doc.add_paragraph()
    add_border_to_paragraph(p, "4472C4")
    p.add_run("Principe : ").bold = True
    p.add_run(
        "Un fichier unique (MEMORY_STATE.md) sert de source de verite pour tous les agents. "
        "AVANT de commencer, chaque agent LIT ce fichier pour comprendre l'etat actuel. "
        "APRES avoir termine, chaque agent MET A JOUR sa section."
    )

    p = doc.add_paragraph()
    p.add_run("Ce fichier contient :").bold = True

    bullets = [
        "Compteurs cumules (46 articles, 22 formules, 18 techniques, 8 axes...)",
        "Version de chaque agent (ce qu'il a produit, ce qui est en attente)",
        "Regles incrementales (AJOUTER vs REECRIRE vs IGNORER)",
        "Registre de feedback utilisateur (pour le MATHTEACHER)",
        "Metriques de qualite tracees entre les executions",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # 3.2 Gate Protocol
    doc.add_heading("3.2 Protocole de Portes (Gates)", level=2)

    p = doc.add_paragraph()
    add_border_to_paragraph(p, "27AE60")
    p.add_run("Principe : ").bold = True
    p.add_run(
        "Aucun agent ne demarre tant que ses prerequis ne sont pas valides. "
        "Le LIBRARIAN attend que les 4 analystes aient termine. "
        "Le SCIENTIST attend que le LIBRARIAN ait cree les index. "
        "Le CHUNKER attend que TOUT soit fait."
    )

    p = doc.add_paragraph()
    p.add_run(
        "Ce protocole empeche les erreurs en cascade : si le COLLECTOR ne trouve que "
        "15 articles (sous le seuil de 20), les analystes recoivent un signal 'PARTIAL' "
        "et adaptent leur travail en consequence."
    )

    # 3.3 DIFF Protocol
    doc.add_heading("3.3 Protocole DIFF (Tracabilite)", level=2)

    p = doc.add_paragraph()
    add_border_to_paragraph(p, "E67E22")
    p.add_run("Principe : ").bold = True
    p.add_run(
        "Chaque agent termine son rapport avec une section DIFF standardisee : "
        "ce qu'il a AJOUTE, MODIFIE, SUPPRIME, et ce qui est INCHANGE. "
        "Cela permet une auditabilite complete — essentiel pour une these doctorale."
    )

    # Example DIFF
    p = doc.add_paragraph()
    run = p.add_run(
        "Exemple DIFF (ANALYST, RUN-002) :\n"
        "  Ajoute : P036-P046 (11 analyses)\n"
        "  Modifie : aucun\n"
        "  Supprime : aucun\n"
        "  Inchange : P001-P035 (35 analyses)"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ========================================================================
    # SECTION 4: CONSTANCY & ITERATION
    # ========================================================================
    doc.add_heading("4. Constance et Repetition : Le Mode Incremental", level=1)

    p = doc.add_paragraph()
    p.add_run("Le probleme classique des agents IA : ").bold = True
    p.add_run(
        "a chaque nouvelle session, l'agent repart de zero. Il recree ce qui existait deja, "
        "perd le contexte precedent, et produit des resultats inconsistants."
    )

    p = doc.add_paragraph()
    p.add_run("Notre solution : ").bold = True
    p.add_run(
        "un systeme de memoire inter-sessions qui transforme chaque execution en amelioration "
        "incrementale. L'agent ne recree pas — il enrichit."
    )

    # Incremental behavior table
    doc.add_heading("Comportement CREATE vs UPDATE par Agent", level=2)

    inc_table = doc.add_table(rows=10, cols=3)
    inc_table.style = 'Medium Shading 1 Accent 1'
    inc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    inc_headers = ['Agent', 'Premiere Execution (CREATE)', 'Executions Suivantes (UPDATE)']
    for i, h in enumerate(inc_headers):
        inc_table.rows[0].cells[i].text = h
        for p in inc_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    inc_data = [
        ["COLLECTOR", "Recherche toutes les annees", "Recherche depuis last_search_date uniquement"],
        ["ANALYST", "Analyse tous les articles", "Analyse uniquement papers_pending"],
        ["MATHEUX", "Extrait toutes les formules", "AJOUTE nouvelles formules au glossaire existant"],
        ["CYBERSEC", "Modelise toutes les menaces", "FUSIONNE nouveaux articles dans l'analyse existante"],
        ["WHITEHACKER", "Extrait toutes les techniques", "AJOUTE techniques T19+, PoC E13+"],
        ["LIBRARIAN", "Construit le filesystem complet", "AJOUTE nouveaux articles, MET A JOUR les index"],
        ["MATHTEACHER", "Cree 7 modules de cours", "AMELIORE (exercices, explications, feedback)"],
        ["SCIENTIST", "Cree 8 axes de recherche", "MET A JOUR axes + scores de confiance"],
        ["CHUNKER", "Chunk tout", "Chunk uniquement fichiers nouveaux/modifies"],
    ]

    for i, row_data in enumerate(inc_data):
        for j, cell_text in enumerate(row_data):
            inc_table.rows[i+1].cells[j].text = cell_text
            for p in inc_table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph("")

    # Callout box
    p = doc.add_paragraph()
    add_border_to_paragraph(p, "8E44AD")
    p.add_run("Exemple concret : ").bold = True
    p.add_run(
        "Le MATHTEACHER a cree 7 modules de cours mathematiques lors du RUN-001. "
        "Lors du RUN-002 (incremental), il ne recree pas les 7 modules — il les AMELIORE : "
        "ajoute des exercices bases sur les nouvelles formules de 2026, raffine les explications "
        "en fonction du feedback utilisateur, et met a jour les sections 'Ou c'est utilise ?' "
        "avec les nouveaux articles. Le cours converge vers la comprehension de l'etudiant."
    )

    # ========================================================================
    # SECTION 5: FEEDBACK LOOP
    # ========================================================================
    doc.add_heading("5. La Boucle de Feedback : L'Agent qui Apprend de Vous", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Le MATHTEACHER illustre parfaitement la puissance de l'iteration agent-humain. "
        "Voici le cycle :"
    )

    # Feedback loop diagram
    feedback_table = doc.add_table(rows=5, cols=2)
    feedback_table.style = 'Light Grid Accent 5'
    feedback_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    feedback_data = [
        ["Etape", "Action"],
        ["1. Generation", "MATHTEACHER produit Module_03 (Theorie de l'Information)"],
        ["2. Lecture", "L'etudiant lit le module et signale : 'Je ne comprends pas la divergence KL'"],
        ["3. Enregistrement", "Le feedback est enregistre dans MEMORY_STATE.md (User Feedback Registry)"],
        ["4. Amelioration", "Au prochain RUN, MATHTEACHER lit le registre et ajoute : nouvelle analogie, "
                           "exemple numerique supplementaire, exercice intermediaire, lien vers prerequis"],
    ]

    for i, (col1, col2) in enumerate(feedback_data):
        feedback_table.rows[i].cells[0].text = col1
        feedback_table.rows[i].cells[1].text = col2
        for p in feedback_table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.font.bold = True
        for p in feedback_table.rows[i].cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.font.bold = True

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run(
        "Ce mecanisme transforme un outil de generation ponctuel en un tuteur adaptatif "
        "qui s'ameliore a chaque interaction. Le cours ne converge pas vers un standard "
        "generique — il converge vers VOTRE niveau de comprehension."
    )

    # ========================================================================
    # SECTION 6: GAIN DE TEMPS
    # ========================================================================
    doc.add_heading("6. Le Gain de Temps : Chiffres Reels", level=1)

    p = doc.add_paragraph()
    p.add_run("Voici la comparaison factuelle entre le travail manuel et le systeme multi-agent :")

    # Time comparison table
    time_table = doc.add_table(rows=11, cols=4)
    time_table.style = 'Medium Shading 1 Accent 1'
    time_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    time_headers = ['Tache', 'Manuel (h)', 'Multi-Agent (min)', 'Gain']
    for i, h in enumerate(time_headers):
        time_table.rows[0].cells[i].text = h
        for p in time_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    time_data = [
        ["Recherche 46 articles", "8-12h", "10 min", "x50"],
        ["34 resumes FR (500 mots)", "34-51h", "8 min", "x250"],
        ["22 formules expliquees", "11-15h", "5 min", "x150"],
        ["Modeles de menaces (34 articles)", "17-25h", "5 min", "x250"],
        ["18 techniques + 12 PoC", "18-24h", "5 min", "x250"],
        ["Organisation filesystem + index", "4-6h", "3 min", "x100"],
        ["7 modules de cours FR", "21-35h", "5 min", "x300"],
        ["8 axes de recherche + SWOT", "8-12h", "3 min", "x200"],
        ["290 chunks RAG", "4-6h", "2 min", "x150"],
        ["TOTAL", "~125-186h", "~45 min", "x170-250"],
    ]

    for i, row_data in enumerate(time_data):
        for j, cell_text in enumerate(row_data):
            time_table.rows[i+1].cells[j].text = cell_text
            for p in time_table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if i == len(time_data) - 1:  # Total row
                        r.font.bold = True

    doc.add_paragraph("")

    # Callout
    p = doc.add_paragraph()
    add_border_to_paragraph(p, "27AE60")
    p.add_run("Facteur multiplicateur reel : x170 a x250. ").bold = True
    p.add_run(
        "Et ce n'est que le RUN-001 (full_search). Les executions suivantes (mode incremental) "
        "prennent ~15 minutes car les agents ne traitent que les nouveaux articles. "
        "Sur un an de these avec des mises a jour hebdomadaires, le gain cumule est "
        "de l'ordre de 500-800 heures."
    )

    # ========================================================================
    # SECTION 7: AGENTIC LOOP
    # ========================================================================
    doc.add_heading("7. La Boucle Agentique : Chaque Agent Pense Avant d'Agir", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Chaque agent n'est pas un simple script — il suit une boucle cognitive autonome "
        "en 8 etapes, inspiree de la methodologie agentique :"
    )

    # Agentic loop diagram
    loop_text = (
        "OBJECTIVE       Quel est mon objectif specifique ?\n"
        "    |\n"
        "DECOMPOSE       Quelles sous-taches sont necessaires ?\n"
        "    |\n"
        "PLAN            Dans quel ordre les executer ?\n"
        "    |\n"
        "ACT             Executer l'action (WebSearch, Read, Write...)\n"
        "    |\n"
        "OBSERVE         Quel est le resultat de mon action ?\n"
        "    |\n"
        "EVALUATE        Le resultat est-il suffisant ? Criteres atteints ?\n"
        "    |\n"
        "REPLAN          Si non -> ajuster le plan et re-executer\n"
        "    |\n"
        "COMPLETE        Si oui -> produire le livrable final + DIFF\n"
    )

    p = doc.add_paragraph()
    run = p.add_run(loop_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p = doc.add_paragraph()
    p.add_run(
        "Cette boucle garantit que chaque agent est auto-correctif. Si le COLLECTOR "
        "ne trouve que 15 articles (sous le seuil de 20), il REPLANIFIE automatiquement "
        "en elargissant ses requetes de recherche. Si l'ANALYST produit un resume "
        "trop court (< 400 mots), il complete avant de passer au suivant."
    )

    # ========================================================================
    # SECTION 8: MODES
    # ========================================================================
    doc.add_heading("8. Six Modes pour Six Besoins", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Le systeme offre 6 modes d'execution, chacun adapte a un besoin specifique :"
    )

    modes_table = doc.add_table(rows=7, cols=3)
    modes_table.style = 'Light Grid Accent 1'
    modes_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    modes_headers = ['Mode', 'Declencheur', 'Agents Actifs']
    for i, h in enumerate(modes_headers):
        modes_table.rows[0].cells[i].text = h
        for p in modes_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    modes_data = [
        ["full_search", '"full", "complete"', "9 agents (pipeline complet)"],
        ["incremental", '"update", "weekly"', "9 agents (nouveaux articles uniquement)"],
        ["analyze_only", '"analyze"', "8 agents (sans COLLECTOR)"],
        ["curriculum_update", '"math", "modules"', "MATHEUX + MATHTEACHER uniquement"],
        ["research_axes", '"axes", "synthese"', "SCIENTIST uniquement"],
        ["rag_refresh", '"rag", "chunks"', "CHUNKER uniquement"],
    ]

    for i, row_data in enumerate(modes_data):
        for j, cell_text in enumerate(row_data):
            modes_table.rows[i+1].cells[j].text = cell_text
            for p in modes_table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    add_border_to_paragraph(p, "4472C4")
    p.add_run("Scheduling automatique : ").bold = True
    p.add_run(
        "Le mode incremental peut etre programme en cron (ex: chaque lundi a 9h). "
        "La bibliographie s'enrichit automatiquement chaque semaine sans intervention humaine."
    )

    # ========================================================================
    # SECTION 9: RESULTATS CONCRETS
    # ========================================================================
    doc.add_heading("9. Resultats Concrets : Ce que le Systeme a Produit", level=1)

    results_table = doc.add_table(rows=13, cols=2)
    results_table.style = 'Medium Shading 1 Accent 5'
    results_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    results_data = [
        ["Metrique", "Valeur (apres 2 RUNs)"],
        ["Articles decouverts", "46 (34 phase 1 + 12 phase 2)"],
        ["Articles analyses en FR", "46 (resumes de ~500 mots chacun)"],
        ["Formules mathematiques", "22+ (avec exemples numeriques)"],
        ["Modeles de menaces", "46 articles mappes MITRE ATT&CK"],
        ["Techniques d'attaque", "18+ avec PoC reproductibles"],
        ["Modules de cours FR", "7 modules (45-55h de formation)"],
        ["Exercices avec solutions", "34 exercices progressifs"],
        ["Axes de recherche", "8 axes identifies et documentes"],
        ["Conjectures validees", "6 (C1-C2 validees, C3-C6 en cours)"],
        ["Chunks RAG", "290 (prets pour ChromaDB)"],
        ["Temps total", "~45 min RUN-001 + ~15 min RUN-002"],
    ]

    for i, (col1, col2) in enumerate(results_data):
        results_table.rows[i].cells[0].text = col1
        results_table.rows[i].cells[1].text = col2
        for p in results_table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.font.bold = True
        for p in results_table.rows[i].cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.font.bold = True

    # ========================================================================
    # SECTION 10: LESSONS LEARNED
    # ========================================================================
    doc.add_heading("10. Lecons Apprises : Ce Qui Marche (et Ce Qui Ne Marche Pas)", level=1)

    doc.add_heading("Ce qui fonctionne remarquablement :", level=2)

    lessons_good = [
        ("Specialisation > Generalisation",
         "Un agent specialise en mathematiques produit un glossaire 10x meilleur "
         "qu'un agent generique. La separation des roles est la cle."),
        ("La memoire partagee est transformatrice",
         "Sans MEMORY_STATE.md, chaque execution repart de zero. Avec, "
         "le systeme s'ameliore — exactement comme une equipe humaine."),
        ("Le mode incremental change tout",
         "Passer de 45 min (full) a 15 min (incremental) rend viable "
         "une mise a jour hebdomadaire automatique."),
        ("Le feedback loop humanise l'IA",
         "Quand le MATHTEACHER ameliore un module parce que VOUS avez dit "
         "'je ne comprends pas', ca cree un partenariat, pas un outil."),
    ]

    for title, desc in lessons_good:
        p = doc.add_paragraph()
        p.add_run(title + " : ").bold = True
        p.add_run(desc)

    doc.add_heading("Les pieges a eviter :", level=2)

    lessons_bad = [
        ("Ne pas sous-estimer les content filters",
         "Les agents de securite manipulent des payloads adversariaux — "
         "les filtres de contenu peuvent bloquer des operations legitimes. "
         "Solution : travailler par metadonnees (comptes, IDs) plutot que contenu brut."),
        ("Choisir le bon modele par agent",
         "Un modele leger (Haiku) suffit pour le COLLECTOR et le LIBRARIAN (taches mecaniques). "
         "Les agents analytiques (ANALYST, SCIENTIST, MATHTEACHER) ont besoin d'un modele "
         "avance (Opus) pour la qualite d'analyse."),
        ("La gate P2->P3 est critique",
         "Si un analyste produit un rapport incomplet, le LIBRARIAN propage l'erreur "
         "dans tous les index. La validation a chaque porte est non-negociable."),
    ]

    for title, desc in lessons_bad:
        p = doc.add_paragraph()
        p.add_run(title + " : ").bold = True
        p.add_run(desc)

    # ========================================================================
    # SECTION 11: APPLICABILITY
    # ========================================================================
    doc.add_heading("11. Au-Dela de la These : Applications Universelles", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Ce pattern multi-agent n'est pas specifique a ma these. Il s'applique a tout "
        "domaine necessitant une recherche structuree :"
    )

    applications = [
        "Veille technologique en entreprise (remplacement des newsletters manuelles)",
        "Due diligence M&A (analyse de centaines de documents legaux)",
        "Revue de litterature medicale (meta-analyses systematiques)",
        "Audit de conformite reglementaire (cartographie des exigences)",
        "Intelligence competitive (surveillance de brevets et publications)",
        "Onboarding technique (generation automatique de formations personnalisees)",
    ]
    for app in applications:
        doc.add_paragraph(app, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run(
        "La cle est toujours la meme : decomposer le travail en roles specialises, "
        "les connecter via une memoire partagee, et iterer incrementalement."
    )

    # ========================================================================
    # CONCLUSION
    # ========================================================================
    doc.add_heading("Conclusion : L'Equipe de Recherche du Futur", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "En 2026, le chercheur isole face a sa pile d'articles est un modele revolu. "
        "Les outils agentiques permettent de constituer une equipe de recherche virtuelle "
        "qui ne dort jamais, ne perd jamais le contexte, et s'ameliore a chaque iteration."
    )

    p = doc.add_paragraph()
    p.add_run(
        "Mon systeme a 9 agents n'est pas parfait — c'est un prototype doctoral. "
        "Mais il demontre une realite emergente : "
    )
    run = p.add_run(
        "la recherche scientifique assistee par IA n'est plus une question de 'si' mais de 'comment'."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p = doc.add_paragraph()
    p.add_run(
        "Le 'comment', c'est la specialisation, la communication inter-agents, "
        "la memoire persistante, et l'iteration incrementale. "
        "C'est transformer un outil generatif en partenaire de recherche."
    )

    # CTA
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("---")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Vous utilisez des agents IA dans votre recherche ? "
        "Partagez vos retours d'experience en commentaires."
    )
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("#IA #RechercheDoctorale #MultiAgent #LLM #Securite #PhD #AgentIA #Automatisation")
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.font.size = Pt(10)

    # ========================================================================
    # ANNEXE: Schema technique detaille
    # ========================================================================
    doc.add_page_break()
    doc.add_heading("Annexe : Schema Technique Detaille du Skill", level=1)

    p = doc.add_paragraph()
    p.add_run("Architecture complete du systeme bibliography-maintainer :").italic = True

    # Detailed technical schema
    tech_schema = (
        "=================================================================\n"
        "              SKILL: bibliography-maintainer\n"
        "         9 Agents | 6 Modes | Memoire Inter-Sessions\n"
        "=================================================================\n"
        "\n"
        "MODES D'EXECUTION:\n"
        "+----------------+  +----------------+  +----------------+\n"
        "| full_search    |  | incremental    |  | analyze_only   |\n"
        "| (9 agents,     |  | (9 agents,     |  | (8 agents,     |\n"
        "|  pipeline       |  |  nouveaux       |  |  sans           |\n"
        "|  complet)       |  |  articles       |  |  COLLECTOR)     |\n"
        "+----------------+  |  uniquement)    |  +----------------+\n"
        "                    +----------------+\n"
        "+----------------+  +----------------+  +----------------+\n"
        "| curriculum_upd |  | research_axes  |  | rag_refresh    |\n"
        "| (MATHEUX +     |  | (SCIENTIST     |  | (CHUNKER       |\n"
        "|  MATHTEACHER)  |  |  uniquement)   |  |  uniquement)   |\n"
        "+----------------+  +----------------+  +----------------+\n"
        "\n"
        "PIPELINE (5 PHASES + GATES):\n"
        "\n"
        "  [P1] COLLECTOR -----> papers.json\n"
        "          |      WebSearch x6 queries\n"
        "          | GATE: >=20 papers, dedup OK\n"
        "          v\n"
        "  [P2] ANALYST ---|--- MATHEUX ---|--- CYBERSEC ---|--- WHITEHACKER\n"
        "       Resume FR  |  Formulas    |  Threats      |  PoC Code\n"
        "       500 mots   |  DAG deps    |  MITRE map    |  Techniques\n"
        "       delta-tags |  Glossaire   |  AEGIS 66     |  Playbook\n"
        "          |       |              |               |\n"
        "          | GATE: All 4 complete, reports valid\n"
        "          v\n"
        "  [P3] LIBRARIAN --> MANIFEST.md + INDEX_BY_DELTA.md\n"
        "          |      Filesystem org, dedup, validation\n"
        "          | GATE: Zero orphans, zero duplicates\n"
        "          v\n"
        "  [P4] MATHTEACHER ---|--- SCIENTIST\n"
        "       7 modules FR  |  8 axes recherche\n"
        "       34 exercices  |  6 conjectures\n"
        "       Quiz eval     |  SWOT these\n"
        "          |          |\n"
        "          | GATE: Both complete\n"
        "          v\n"
        "  [P5] CHUNKER --> chunks_for_rag.jsonl (290 chunks)\n"
        "                   ingest_to_chromadb.py\n"
        "\n"
        "MEMOIRE INTER-SESSIONS:\n"
        "\n"
        "  MEMORY_STATE.md <----> Tous les agents (lecture/ecriture)\n"
        "       |                   |\n"
        "       +-- Compteurs       +-- Versions agents\n"
        "       +-- Regles incr.    +-- Feedback user\n"
        "       +-- Metriques       +-- Papers pending\n"
        "       |\n"
        "  EXECUTION_LOG.jsonl <-- Historique (append-only)\n"
        "       |\n"
        "       +-- Un JSON/ligne par RUN\n"
        "       +-- Tendances et regressions\n"
    )

    p = doc.add_paragraph()
    run = p.add_run(tech_schema)
    run.font.name = 'Consolas'
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # ========================================================================
    # ANNEXE 2: Inter-agent data flow
    # ========================================================================
    doc.add_heading("Flux de Donnees Inter-Agents", level=2)

    flow_table = doc.add_table(rows=10, cols=4)
    flow_table.style = 'Light Grid Accent 1'
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    flow_headers = ['Source', 'Destination', 'Donnee Transmise', 'Format']
    for i, h in enumerate(flow_headers):
        flow_table.rows[0].cells[i].text = h
        for p in flow_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(8)

    flow_data = [
        ["COLLECTOR", "ANALYST", "Metadonnees articles", "papers.json (JSONL)"],
        ["COLLECTOR", "MATHEUX", "Articles avec formules", "papers.json"],
        ["ANALYST", "LIBRARIAN", "Analyses completees", "P0XX_analysis.md"],
        ["MATHEUX", "MATHTEACHER", "Glossaire + DAG", "GLOSSAIRE_DETAILED.md"],
        ["MATHEUX", "LIBRARIAN", "Glossaire unifie", "GLOSSAIRE_MATHEMATIQUE.md"],
        ["CYBERSEC", "SCIENTIST", "Matrice menaces", "THREAT_ANALYSIS.md"],
        ["WHITEHACKER", "SCIENTIST", "Playbook attaques", "RED_TEAM_PLAYBOOK.md"],
        ["LIBRARIAN", "SCIENTIST", "Index organises", "MANIFEST.md + INDEX"],
        ["TOUS", "CHUNKER", "Tous les livrables", "*.md -> JSONL chunks"],
    ]

    for i, row_data in enumerate(flow_data):
        for j, cell_text in enumerate(row_data):
            flow_table.rows[i+1].cells[j].text = cell_text
            for p in flow_table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    # Save
    output_path = os.path.join(
        "C:/Users/pizzif/Documents/GitHub/poc_medical/research_archive/_staging",
        "Article_LinkedIn_Multi_Agent_Research_Team.docx"
    )
    doc.save(output_path)
    print(f"Article genere : {output_path}")
    print(f"Nombre de sections : 11 + 2 annexes")
    print(f"Schemas integres : 6 (pipeline, communications, boucle agentique, flux donnees, modes, incremental)")

if __name__ == "__main__":
    create_article()
